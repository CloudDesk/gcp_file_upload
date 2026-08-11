from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "invoice" / "revoinvoiceservice.docx"


def set_fixed_layout(table) -> None:
    table.autofit = False
    table_properties = table._tbl.tblPr
    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_margins(cell, top=35, start=45, bottom=35, end=45) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)

    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        margin = margins.find(qn(f"w:{edge}"))
        if margin is None:
            margin = OxmlElement(f"w:{edge}")
            margins.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def set_cell_text(
    cell,
    text: str,
    *,
    font_size: float,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def configure_table(table, widths: list[float]) -> None:
    while len(table.columns) < len(widths):
        table.add_column(Inches(1))

    if len(table.columns) != len(widths):
        raise RuntimeError(
            f"Unexpected column count: found {len(table.columns)}, expected {len(widths)}"
        )

    set_fixed_layout(table)
    grid_columns = table._tbl.tblGrid.gridCol_lst
    for column_index, width_inches in enumerate(widths):
        width = Inches(width_inches)
        table.columns[column_index].width = width
        if column_index < len(grid_columns):
            grid_columns[column_index].set(qn("w:w"), str(width.emu))
        for row in table.rows:
            cell = row.cells[column_index]
            cell.width = width
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(int(width.twips)))
            cell_width.set(qn("w:type"), "dxa")


def remove_extra_blank_row(table) -> None:
    if len(table.rows) <= 2:
        return

    extra_row = table.rows[2]
    if any(cell.text.strip() for cell in extra_row.cells):
        raise RuntimeError("Expected the third table row to be blank")
    table._tbl.remove(extra_row._tr)


def update_product_table(table) -> None:
    widths = [0.40, 1.75, 1.55, 0.55, 0.85, 1.20, 1.55]
    headers = [
        "S.NO",
        "PRODUCT NAME",
        "DESCRIPTION",
        "QTY",
        "HSN CODE",
        "UNIT PRICE",
        "TOTAL AMOUNT",
    ]
    placeholders = [
        "{#invoicedata}{#items}{id}",
        "{productname}",
        "{description}",
        "{quantity}",
        "{hsncode}",
        "{price}",
        "{totalamount}{/}{/}",
    ]

    remove_extra_blank_row(table)
    configure_table(table, widths)
    for column_index, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[column_index],
            header,
            font_size=8,
            bold=True,
        )
        set_cell_text(
            table.rows[1].cells[column_index],
            placeholders[column_index],
            font_size=8.5,
            alignment=(
                WD_ALIGN_PARAGRAPH.LEFT
                if column_index in (1, 2)
                else WD_ALIGN_PARAGRAPH.CENTER
            ),
        )


def update_service_table(table) -> None:
    widths = [0.45, 2.60, 1.20, 1.55, 2.05]
    headers = ["S.NO", "DESCRIPTION", "SAC CODE", "UNIT PRICE", "TOTAL AMOUNT"]
    placeholders = [
        "{#servicedata}{#items}{id}",
        "{description}",
        "{saccode}",
        "{price}",
        "{totalamount}{/}{/}",
    ]

    remove_extra_blank_row(table)
    configure_table(table, widths)
    for column_index, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[column_index],
            header,
            font_size=9,
            bold=True,
        )
        set_cell_text(
            table.rows[1].cells[column_index],
            placeholders[column_index],
            font_size=9,
            alignment=(
                WD_ALIGN_PARAGRAPH.LEFT
                if column_index == 1
                else WD_ALIGN_PARAGRAPH.CENTER
            ),
        )


def update_tax_labels(document) -> None:
    replacements = {
        "Tax%: {#invoicedata}{tax}{/}": "Tax: {#invoicedata}{taxlabel}{/}",
        "Tax%: {#servicedata}{tax}{/}": "Tax: {#servicedata}{taxlabel}{/}",
    }
    for paragraph in document.paragraphs:
        replacement = replacements.get(paragraph.text.strip())
        if replacement is None:
            continue

        for run in list(paragraph.runs):
            paragraph._p.remove(run._element)
        run = paragraph.add_run(replacement)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")


def remove_service_type_line(document) -> None:
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip().startswith("Service Type:"):
            paragraph._element.getparent().remove(paragraph._element)


def main() -> None:
    document = Document(TEMPLATE_PATH)
    if len(document.tables) < 3:
        raise RuntimeError("Service invoice template does not contain both item tables")

    update_product_table(document.tables[1])
    update_service_table(document.tables[2])
    update_tax_labels(document)
    remove_service_type_line(document)
    document.save(TEMPLATE_PATH)
    print(f"Updated {TEMPLATE_PATH}")


if __name__ == "__main__":
    main()
