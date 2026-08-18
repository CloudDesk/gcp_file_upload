from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "invoice" / "revoinvoiceproduct.docx"

REPLACEMENTS = {
    "{customername}": "{shippingcustomername}",
    "{customeraddress}": "{shippingcustomeraddress}",
    "Phone Number: +91 {customerphonenumber}": "Phone Number: +91 {shippingcustomerphonenumber}",
    "GST: {customergstnumber}": "GST: {shippingcustomergstnumber}",
}


def main() -> None:
    document = Document(TEMPLATE_PATH)
    shipping_started = False

    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "SHIPPING TO":
            shipping_started = True
            continue
        if not shipping_started:
            continue

        replacement = REPLACEMENTS.get(paragraph.text.strip())
        if replacement is None:
            continue

        for run in list(paragraph.runs):
            paragraph._p.remove(run._element)
        paragraph.add_run(replacement)

    document.save(TEMPLATE_PATH)
    print(f"Updated {TEMPLATE_PATH}")


if __name__ == "__main__":
    main()
